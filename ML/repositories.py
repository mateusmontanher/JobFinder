"""Infrastructure adapters for resume and job text sources."""

from __future__ import annotations

import csv
import json
import os
from collections.abc import Iterator
from io import BytesIO
from pathlib import Path
from typing import Any, Protocol

from docx import Document


class ResumeRepository(Protocol):
    def load_text(self) -> str: ...


def document_text(data: bytes) -> str:
    document = Document(BytesIO(data))
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        parts.extend(cell.text for row in table.rows for cell in row.cells)
    return "\n".join(part for part in parts if part.strip())


class PostgresResumeRepository:
    """Loads the most recently uploaded resume without retaining a connection."""

    def __init__(self, connect=None) -> None:
        self._connect = connect or self._default_connect

    @staticmethod
    def _default_connect():
        import psycopg2
        return psycopg2.connect(
            database=os.getenv("DB_NAME"), user=os.getenv("DB_USER"),
            password=os.getenv("DB_PASSWORD"), host=os.getenv("DB_HOST"),
            port=os.getenv("DB_PORT"),
        )

    def load_text(self) -> str:
        with self._connect() as connection:
            with connection.cursor() as cursor:
                cursor.execute(
                    "SELECT file_data FROM curriculum ORDER BY uploaded_at DESC, id DESC LIMIT 1"
                )
                row = cursor.fetchone()
        if not row or not row[0]:
            raise ValueError("no resume is stored in the curriculum table")
        return document_text(bytes(row[0]))


class FileResumeRepository:
    def __init__(self, path: Path) -> None:
        self.path = path

    def load_text(self) -> str:
        suffix = self.path.suffix.casefold()
        if suffix == ".docx":
            return document_text(self.path.read_bytes())
        if suffix == ".txt":
            return self.path.read_text(encoding="utf-8")
        if suffix == ".json":
            payload = json.loads(self.path.read_text(encoding="utf-8"))
            if isinstance(payload, str):
                return payload
            for key in ("resume", "text", "raw_text"):
                if isinstance(payload, dict) and isinstance(payload.get(key), str):
                    return payload[key]
        if suffix == ".csv":
            with self.path.open(encoding="utf-8-sig", newline="") as source:
                row = next(csv.DictReader(source), None)
            if row:
                for key in ("resume", "text", "raw_text"):
                    if row.get(key):
                        return row[key]
        raise ValueError(f"unsupported or invalid resume file: {self.path}")


def iter_job_records(path: Path) -> Iterator[dict[str, Any]]:
    """Read jobs from JSON, JSONL, CSV, TXT, or a folder of those formats."""
    if path.is_dir():
        for child in sorted(path.iterdir()):
            if child.suffix.casefold() in {".json", ".jsonl", ".csv", ".txt"}:
                yield from iter_job_records(child)
        return
    suffix = path.suffix.casefold()
    if suffix == ".csv":
        with path.open(encoding="utf-8-sig", newline="") as source:
            yield from csv.DictReader(source)
        return
    if suffix == ".txt":
        yield {"id": path.stem, "title": path.stem, "description": path.read_text(encoding="utf-8")}
        return
    text = path.read_text(encoding="utf-8")
    payloads = [json.loads(line) for line in text.splitlines() if line.strip()] if suffix == ".jsonl" else json.loads(text)
    if isinstance(payloads, dict):
        payloads = payloads.get("jobs", [payloads])
    if not isinstance(payloads, list):
        raise ValueError(f"job input must contain records: {path}")
    for payload in payloads:
        if isinstance(payload, dict):
            yield payload


def job_fields(record: dict[str, Any], index: int = 0) -> tuple[str, str, str]:
    title = next((record[key].strip() for key in ("title", "job_title", "name") if isinstance(record.get(key), str)), "")
    description = next((record[key].strip() for key in ("description", "job_text", "summary") if isinstance(record.get(key), str)), "")
    identifier = next((str(record[key]) for key in ("id", "job_id", "url", "link") if record.get(key)), f"job-{index}")
    return identifier, title, description
